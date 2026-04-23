"""
Sports Leaderboard — single-process Streamlit app.

No Flask backend. State is shared across all browser sessions via
@st.cache_resource (all phones hit the same Python process).
Schedule + scores are persisted to a JSON file so data survives restarts.

Gemini API key is read from config.yaml (gemini.api_key).
"""

import base64
import copy
import json
import logging
import os
import threading
from pathlib import Path
from typing import Dict, List, Optional

import random
import time as _time

import requests

import importlib
import quips as _quips_module
importlib.reload(_quips_module)
_ALL_QUIPS = _quips_module.QUIPS

import pandas as pd
import streamlit as st
import yaml

logging.basicConfig(level=logging.INFO)

# ===========================================================================
# Config  (config.yaml)
# ===========================================================================

@st.cache_resource
def _cfg() -> dict:
    """
    Load config in priority order:
      1. st.secrets  — Streamlit Community Cloud (secrets set in the dashboard)
      2. config.yaml — local development
    """
    # ── Streamlit Community Cloud ──────────────────────────────────────────
    try:
        if "gemini" in st.secrets:
            return {
                "gemini": dict(st.secrets["gemini"]),
                "app":    dict(st.secrets.get("app", {})),
            }
    except Exception:
        pass

    # ── Local dev: config.yaml ─────────────────────────────────────────────
    p = Path("config.yaml")
    if p.exists():
        try:
            with open(p, encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
        except Exception as e:
            logging.warning("Could not read config.yaml: %s", e)
    return {}


def _gemini_key() -> str:
    key = _cfg().get("gemini", {}).get("api_key", "")
    if key and not key.startswith("YOUR_"):
        return key
    return ""


def _admin_password() -> str:
    try:
        return st.secrets["app"]["admin_password"]
    except Exception:
        return _cfg().get("app", {}).get("admin_password", "kaushik28")




def _data_dir() -> Path:
    # Explicit env var wins (useful for Azure App Service)
    env = os.getenv("DATA_DIR")
    if env:
        return Path(env)
    # config / secrets value
    cfg_dir = _cfg().get("app", {}).get("data_dir", "")
    if cfg_dir:
        return Path(cfg_dir)
    # Default: ./data locally, /tmp on Linux-based clouds
    import platform
    return Path("./data") if platform.system() == "Windows" else Path("/tmp/leaderboard")


def _default_rounds() -> int:
    return int(_cfg().get("app", {}).get("default_rounds", 12))


# ===========================================================================
# Shared persistent state
# ===========================================================================

_file_lock = threading.Lock()

# ---------------------------------------------------------------------------
# GitHub persistence helpers
# ---------------------------------------------------------------------------

def _gh_cfg() -> dict:
    """Return the [github] config block from secrets or config.yaml."""
    try:
        if "github" in st.secrets:
            return dict(st.secrets["github"])
    except Exception:
        pass
    return _cfg().get("github", {})


def _gh_headers() -> dict:
    token = _gh_cfg().get("token", "")
    return {"Authorization": f"token {token}",
            "Accept": "application/vnd.github.v3+json"}


def _gh_file_url(gh_path: str) -> str:
    repo = _gh_cfg().get("repo", "")
    return f"https://api.github.com/repos/{repo}/contents/{gh_path}"


# ── Generic GitHub file helpers ──────────────────────────────────────────────

def _github_load_file(gh_path: str, default: Optional[dict] = None) -> Optional[dict]:
    """Fetch a JSON file from GitHub. Returns default if 404, None on error."""
    if not _gh_cfg().get("token") or not _gh_cfg().get("repo"):
        return None
    try:
        r = requests.get(_gh_file_url(gh_path), headers=_gh_headers(), timeout=10)
        if r.status_code == 200:
            content = base64.b64decode(r.json()["content"]).decode("utf-8")
            return json.loads(content)
        if r.status_code == 404:
            return default
        logging.warning("GitHub load HTTP %s (%s)", r.status_code, gh_path)
    except Exception as exc:
        logging.warning("GitHub load error (%s): %s", gh_path, exc)
    return None


def _github_save_file(gh_path: str, data: dict, commit_msg: str = "update") -> None:
    """Push a JSON file to GitHub (create or update)."""
    if not _gh_cfg().get("token") or not _gh_cfg().get("repo"):
        return
    try:
        hdrs = _gh_headers()
        url  = _gh_file_url(gh_path)
        sha  = ""
        r_get = requests.get(url, headers=hdrs, timeout=10)
        if r_get.status_code == 200:
            sha = r_get.json().get("sha", "")
        content_b64 = base64.b64encode(
            json.dumps(data, indent=2).encode("utf-8")
        ).decode("utf-8")
        payload: dict = {"message": commit_msg, "content": content_b64}
        if sha:
            payload["sha"] = sha
        r = requests.put(url, headers=hdrs, json=payload, timeout=15)
        if r.status_code in (200, 201):
            logging.info("GitHub save OK (%s)", gh_path)
        else:
            logging.warning("GitHub save HTTP %s (%s): %s",
                            r.status_code, gh_path, r.text[:200])
    except Exception as exc:
        logging.warning("GitHub save error (%s): %s", gh_path, exc)


# ── Convenience wrappers for session.json ────────────────────────────────────

def _github_load() -> Optional[dict]:
    return _github_load_file("data/session.json", default=_empty_state())


def _publish_leaderboard(lb: list) -> None:
    """Append current leaderboard snapshot with timestamp to published_results.json on GitHub."""
    from datetime import datetime, timezone
    existing = _github_load_file("data/published_results.json", default={"results": []})
    if not isinstance(existing, dict):
        existing = {"results": []}
    _state = _get()
    entry = {
        "timestamp": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC"),
        "girl_names": list(_state.get("girl_names", [])),
        "players": [
            {
                "name":       p["name"],
                "wins":       p["games_won"],
                "losses":     p["games_lost"],
                "net_points": p["net_points"],
            }
            for p in lb
        ],
    }
    existing["results"].append(entry)
    _github_save_file("data/published_results.json", existing, "leaderboard: publish results")

def _github_save(state: dict) -> None:
    _github_save_file("data/session.json", state, "leaderboard: update session")


# ── Users / phone-number access list ─────────────────────────────────────────

def _default_users() -> dict:
    return {"allowed_phones": []}

@st.cache_resource
def _users_box() -> dict:
    gh = _github_load_file("data/users.json", default=_default_users())
    return {"u": gh if gh is not None else _default_users()}

def _get_users() -> dict:
    return _users_box()["u"]

def _put_users(users: dict) -> None:
    _users_box()["u"] = users
    _github_save_file("data/users.json", users, "leaderboard: update users")


# ---------------------------------------------------------------------------
# Shared persistent state
# ---------------------------------------------------------------------------

def _data_file() -> Path:
    return _data_dir() / "session.json"


def _empty_state() -> dict:
    return {
        "players": [],
        "skill_levels": {},
        "schedule": [],
        "scores": {},
        "session_active": False,
        "critics_choice": None,
    }


def _load_state() -> dict:
    """Load state: GitHub first (persistent), local file as fallback."""
    gh = _github_load()
    if gh is not None:
        logging.info("State loaded from GitHub")
        return gh
    # fallback: local JSON file
    try:
        f = _data_file()
        if f.exists():
            with open(f, encoding="utf-8") as fp:
                logging.info("State loaded from local file")
                return json.load(fp)
    except Exception as e:
        logging.warning("Could not load local session file: %s", e)
    return _empty_state()


@st.cache_resource
def _box() -> dict:
    """Singleton dict shared across ALL browser sessions."""
    return {"s": _load_state()}


def _get() -> dict:
    return _box()["s"]


def _put(state: dict) -> None:
    """Update shared state, persist to GitHub and local file."""
    with _file_lock:
        _box()["s"] = state
        # GitHub (primary — survives restarts / redeploys)
        _github_save(state)
        # Local file (fast fallback)
        try:
            d = _data_dir()
            d.mkdir(parents=True, exist_ok=True)
            tmp = _data_file().with_suffix(".tmp")
            with open(tmp, "w", encoding="utf-8") as f:
                json.dump(state, f, indent=2)
            tmp.replace(_data_file())
        except Exception as e:
            logging.warning("Could not persist state locally: %s", e)


# ===========================================================================
# Page config + mobile CSS
# ===========================================================================

st.set_page_config(
    page_title="Sports Leaderboard",
    page_icon="🎾",
    layout="centered",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <style>
    /* ── Layout ─────────────────────────────────────────── */
    .main .block-container {
        padding: 1rem 0.75rem 5rem 0.75rem;
        max-width: 520px;
    }

    /* ── Page header banner ──────────────────────────────── */
    .page-header {
        background: linear-gradient(135deg, #1565C0 0%, #0288D1 100%);
        padding: 1rem 1.15rem 0.9rem;
        border-radius: 14px;
        margin-bottom: 1.1rem;
        color: white;
    }
    .page-header h1 {
        margin: 0;
        font-size: 1.45rem;
        font-weight: 700;
        line-height: 1.2;
        color: white;
    }
    .page-header p {
        margin: 0.2rem 0 0;
        font-size: 0.82rem;
        opacity: 0.88;
        color: white;
    }

    /* ── Section label ───────────────────────────────────── */
    .section-label {
        font-size: 0.75rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        color: #888;
        margin: 1.1rem 0 0.35rem;
    }

    /* ── Team cards (court page) ─────────────────────────── */
    .team-card-a {
        background: #0D2137;
        border-left: 4px solid #29B6F6;
        padding: 0.55rem 0.75rem;
        border-radius: 8px;
        margin-bottom: 0.4rem;
        font-size: 0.95rem;
        line-height: 1.5;
        color: #E8EAF0;
    }
    .team-card-b {
        background: #2A0D12;
        border-left: 4px solid #EF5350;
        padding: 0.55rem 0.75rem;
        border-radius: 8px;
        margin-bottom: 0.6rem;
        font-size: 0.95rem;
        line-height: 1.5;
        color: #E8EAF0;
    }

    /* ── Buttons ─────────────────────────────────────────── */
    .stButton > button {
        min-height: 48px;
        font-size: 1rem;
        border-radius: 10px;
        width: 100%;
        font-weight: 500;
        transition: opacity 0.15s;
    }
    .stButton > button:active { opacity: 0.82; }

    /* ── Number / text inputs ────────────────────────────── */
    .stNumberInput input {
        font-size: 1.35rem !important;
        height: 52px !important;
        text-align: center !important;
    }
    .stNumberInput [data-testid="stNumberInputStepDown"],
    .stNumberInput [data-testid="stNumberInputStepUp"] {
        width: 40px; height: 52px;
    }
    .stTextInput input  { font-size: 1rem !important; height: 44px !important; }
    .stSelectbox > div > div { font-size: 1rem !important; min-height: 44px; }

    /* ── Win buttons (court page) ───────────────────────── */
    [data-testid="stHorizontalBlock"]:has(.team-card-a) + div .stButton > button,
    [data-testid="stHorizontalBlock"]:has(.team-card-b) + div .stButton > button {
        min-height: 44px !important;
        font-size: 0.88rem !important;
        padding: 0.25rem 0.3rem !important;
    }


    /* ── Roster radio buttons ────────────────────────────── */
    div[data-testid="stRadio"] > div {
        gap: 0.4rem;
        padding-top: 0.45rem;
    }
    div[data-testid="stRadio"] label {
        font-size: 0.82rem !important;
        padding: 0.2rem 0.1rem;
    }

    /* ── Expanders ───────────────────────────────────────── */
    details { border-radius: 10px !important; }
    details summary {
        font-size: 1rem;
        padding: 0.65rem 0;
        line-height: 1.4;
        font-weight: 500;
    }

    /* ── Metrics ─────────────────────────────────────────── */
    [data-testid="stMetric"] {
        background: #1A1F2E;
        border-radius: 10px;
        padding: 0.5rem 0.4rem;
        text-align: center;
    }
    [data-testid="stMetricDelta"] { font-size: 0.75rem; }
    [data-testid="stMetricLabel"] { font-size: 0.78rem !important; }
    [data-testid="stMetricValue"] { font-size: 1.3rem !important; }

    /* ── Progress bar ────────────────────────────────────── */
    [data-testid="stProgressBar"] { height: 8px; border-radius: 5px; }
    [data-testid="stProgressBar"] > div > div {
        background: linear-gradient(90deg, #1565C0, #0288D1) !important;
    }

    /* ── DataFrame ───────────────────────────────────────── */
    [data-testid="stDataFrame"] { font-size: 0.85rem; border-radius: 8px; overflow: hidden; }

    /* ── Tabs ────────────────────────────────────────────── */
    .stTabs [data-baseweb="tab-list"] { gap: 4px; }
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px 8px 0 0;
        padding: 0.4rem 0.9rem;
        font-size: 0.9rem;
        font-weight: 500;
    }

    /* ── Alert / info boxes ──────────────────────────────── */
    [data-testid="stAlert"] { border-radius: 10px; }

    /* ── Boys / Girls section headings ──────────────────── */
    .roster-heading-boys { color: #5B9BD5; font-weight: 700; font-size: 1rem; margin: 0.6rem 0 0.25rem 0; }
    .roster-heading-girls { color: #E07090; font-weight: 700; font-size: 1rem; margin: 0.6rem 0 0.25rem 0; }

    /* ── Hide Streamlit chrome ───────────────────────────── */
    #MainMenu, footer { visibility: hidden; }
    header[data-testid="stHeader"] { height: 2rem; }
    </style>
    """,
    unsafe_allow_html=True,
)


# ===========================================================================
# Session-state init  (per-browser UI state, not shared)
# ===========================================================================

def _init_ui():
    defaults = {
        "page": "setup",
        "num_boys":    6,
        "num_girls":   4,
        "boy_names":   [""] * 6,
        "girl_names":  [""] * 4,
        "num_players": 10,
        "num_courts": 2,
        "games_per_hour": 5,
        "player_names": [""] * 10,
        "skill_visible": False,
        "show_skill_pw": False,
        "show_gen_pw": False,
        "show_reset_pw": False,
        "phone_verified": False,
        "verified_phone": "",
        "show_admin_pw": False,
        "show_admin_panel": False,
        "special_instructions": "Avoid a team with 2 girls if the opponent team has a boy.",
        "phone_add_counter": 0,
        "show_pub_pw": False,
        "show_rst_pw": False,
        "show_rst1_pw": False,
        "show_rst1_select": False,
        "show_rst_confirm": False,
        "show_rst1_confirm": False,
        "rst1_pending_delete": [],
        "_active_setup_tab": 0,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # Restore court hours from persisted state on every fresh browser load.
    # This must run in _init_ui (not just show_setup) so the slider values are
    # correct regardless of which page the user lands on first.
    if not st.session_state.get("_court_hours_init"):
        _s = _get()
        if _s.get("session_active") and _s.get("court_hours"):
            for _c, _hrs in _s["court_hours"].items():
                st.session_state[f"court_hours_{_c}"] = float(_hrs)
        st.session_state._court_hours_init = True


_init_ui()

# ── Auto-login: query-param check (same tab / bookmarked URL) ────────────────
if not st.session_state.phone_verified:
    _qp_phone = st.query_params.get("lb_auth", "")
    if _qp_phone:
        _cu = _get_users()
        if _qp_phone in _cu.get("allowed_phones", []):
            st.session_state.phone_verified = True
            st.session_state.verified_phone = _qp_phone
            st.rerun()
    else:
        # No query param — try localStorage (cross-tab, new session same browser)
        st.components.v1.html(
            """<script>
            (function(){
                try {
                    var d = JSON.parse(window.localStorage.getItem("lb_auth") || "null");
                    if (d && d.p && d.t && (Date.now()/1000 - d.t) < 7200) {
                        var u = new URL(window.top.location.href);
                        if (!u.searchParams.get("lb_auth")) {
                            u.searchParams.set("lb_auth", d.p);
                            window.top.location.replace(u.toString());
                        }
                    } else if (d) {
                        window.localStorage.removeItem("lb_auth");
                    }
                } catch(e) {}
            })();
            </script>""",
            height=0,
        )

# ===========================================================================
# Sidebar
# ===========================================================================

with st.sidebar:
    st.markdown("## 🏸 Sports Leaderboard")
    st.divider()

    _s = _get()
    _num_courts_state = _s.get("num_courts", st.session_state.get("num_courts", 2))

    _sb_nav = [("setup", "⚙️  Setup & Schedule")]
    for _c in range(1, _num_courts_state + 1):
        _sb_nav.append((f"court{_c}", f"🏟️  Court {_c}"))
    _sb_nav.append(("leaderboard", "🏆  Leaderboard"))
    _sb_nav.append(("alltime", "📊  All Time"))

    for pid, label in _sb_nav:
        if st.button(
            label,
            key=f"sb_{pid}",
            type="primary" if st.session_state.page == pid else "secondary",
            use_container_width=True,
        ):
            st.session_state.page = pid
            st.rerun()

    st.divider()
    if _s.get("session_active"):
        st.success(f"✅ {len(_s['players'])} players · {_num_courts_state} courts")
        total = len(_s["scores"])
        done  = sum(1 for v in _s["scores"].values() if v.get("submitted"))
        if total:
            st.progress(done / total, text=f"{done}/{total} games done")
    else:
        st.info("No active session")

    if not _gemini_key():
        st.warning("⚠️ Gemini key missing.\nAdd it to config.yaml.")


# ===========================================================================
# Bottom nav bar  (same 4 buttons on every page)
# ===========================================================================

def _nav(active: str) -> None:
    num_courts = _get().get("num_courts", st.session_state.get("num_courts", 2))
    nav = [("setup", "⚙️", "Setup")]
    for c in range(1, num_courts + 1):
        nav.append((f"court{c}", "🏟", f"Crt {c}"))
    nav.append(("leaderboard", "🏆", "Leaderboard"))
    nav.append(("alltime", "📊", "All Time"))

    cols = st.columns(len(nav))
    for col, (pid, icon, label) in zip(cols, nav):
        with col:
            if st.button(
                f"{icon}\n{label}",
                key=f"bn_{pid}",
                type="primary" if active == pid else "secondary",
                use_container_width=True,
            ):
                st.session_state.page = pid
                st.rerun()

    # Normalize button heights: the 🏆 emoji renders taller than others in some browsers
    st.components.v1.html(
        """<script>
        (function() {
            function fixNavBtns() {
                var main = window.parent.document.querySelector('[data-testid="stMain"]');
                if (!main) return;
                var hblocks = main.querySelectorAll('[data-testid="stHorizontalBlock"]');
                if (!hblocks.length) return;
                var navBlock = hblocks[0];
                navBlock.querySelectorAll('button').forEach(function(b) {
                    b.style.setProperty('height', '3.2rem', 'important');
                    b.style.setProperty('max-height', '3.2rem', 'important');
                    b.style.setProperty('overflow', 'hidden', 'important');
                    b.style.setProperty('font-size', '0.7rem', 'important');
                    b.style.setProperty('line-height', '1.15', 'important');
                    b.style.setProperty('padding', '0.1rem 0.15rem', 'important');
                });
            }
            setTimeout(fixNavBtns, 100);
            setTimeout(fixNavBtns, 400);
        })();
        </script>""",
        height=0,
    )


# ===========================================================================
# Page: Setup & Schedule
# ===========================================================================

def show_setup() -> None:
    # Restore saved session values only once per browser session (on first load).
    # Skipping on subsequent reruns prevents the saved state from overwriting
    # widget changes the user is actively making.
    _s = _get()
    if _s.get("session_active") and not st.session_state.get("_setup_restored"):
        st.session_state.num_courts     = _s.get("num_courts",     st.session_state.get("num_courts", 2))
        st.session_state.games_per_hour = _s.get("games_per_hour", st.session_state.get("games_per_hour", 5))
        for _c, _hrs in _s.get("court_hours", {}).items():
            st.session_state[f"court_hours_{_c}"] = float(_hrs)
        if _s.get("players"):
            _all_p      = list(_s["players"])
            _saved_gset = set(_s.get("girl_names", []))
            _bnames     = [p for p in _all_p if p not in _saved_gset]
            _gnames     = [p for p in _all_p if p in _saved_gset]
            st.session_state.num_boys      = len(_bnames)
            st.session_state.num_girls     = len(_gnames)
            st.session_state.boy_names     = _bnames
            st.session_state.girl_names    = _gnames
            st.session_state.num_players   = len(_all_p)
            st.session_state.player_names  = _all_p
            for _i, _nm in enumerate(_bnames):
                st.session_state[f"pname_boy_{_i}"] = _nm
            for _i, _nm in enumerate(_gnames):
                st.session_state[f"pname_girl_{_i}"] = _nm
        if "special_instructions" in _s:
            st.session_state.special_instructions = _s["special_instructions"]
        st.session_state._setup_restored = True

    st.markdown(
        '<div class="page-header">'
        '<h1>⚙️ Setup</h1>'
        '<p>Configure players, courts &amp; generate your schedule</p>'
        '</div>',
        unsafe_allow_html=True,
    )

    tab_p, tab_s = st.tabs(["👥 Players & Courts", "⚙️ Generate Schedule"])

    # Re-select the Generate Schedule tab after rerun if user was on it
    if st.session_state.get("_active_setup_tab") == 1:
        st.session_state._active_setup_tab = 0
        st.components.v1.html(
            """<script>
            setTimeout(function() {
                var tabs = window.parent.document.querySelectorAll('[data-baseweb="tab"]');
                if (tabs && tabs.length > 1) { tabs[1].click(); }
            }, 150);
            </script>""",
            height=0,
        )

    # ── Players tab ──────────────────────────────────────────────────────────
    with tab_p:
        # Number of courts
        num_courts = st.number_input(
            "Number of courts", min_value=1, max_value=4,
            value=st.session_state.get("num_courts", 2), step=1,
        )
        if num_courts != st.session_state.get("num_courts", 2):
            for _c in range(1, st.session_state.get("num_courts", 2) + 1):
                _hk = f"court_hours_{_c}"
                if _hk in st.session_state:
                    st.session_state[f"_bk_{_hk}"] = st.session_state[_hk]
            st.session_state.num_courts = num_courts
            st.rerun()

        # Single rate slider
        games_per_hour = st.slider(
            "Games /per court/ per hour",
            min_value=1, max_value=12, step=1,
            value=st.session_state.get("games_per_hour", 5),
        )
        if games_per_hour != st.session_state.get("games_per_hour", 5):
            st.session_state.games_per_hour = games_per_hour
        mins_per_game = round(60 / games_per_hour, 1)
        st.caption(f"~{mins_per_game} min per game")

        # Per-court hours-booked sliders
        st.markdown('<div class="section-label">Hours booked per court</div>', unsafe_allow_html=True)
        num_games_per_court: Dict[int, int] = {}
        for c in range(1, num_courts + 1):
            hrs_key = f"court_hours_{c}"
            if hrs_key not in st.session_state:
                bk_key = f"_bk_{hrs_key}"
                if bk_key in st.session_state:
                    st.session_state[hrs_key] = st.session_state[bk_key]
                else:
                    _saved_ch = _get().get("court_hours", {})
                    _v = _saved_ch.get(c) or _saved_ch.get(str(c))
                    st.session_state[hrs_key] = float(_v) if _v is not None else 2.0
            court_hrs = st.slider(
                f"Court {c}",
                min_value=0.5, max_value=6.0, step=0.5,
                key=hrs_key,
            )
            num_games_per_court[c] = max(1, round(games_per_hour * court_hrs))
            st.caption(f"→ {num_games_per_court[c]} games")

        # Session summary info box
        n           = st.session_state.get("num_boys", 6) + st.session_state.get("num_girls", 4)
        st.session_state.num_players = n
        total_games = sum(num_games_per_court.values())
        detail      = "  \n".join(
            f"Court {c}: {g} games ({st.session_state.get(f'court_hours_{c}', 2.0):.1f}h)"
            for c, g in num_games_per_court.items()
        )
        avg_games   = round(total_games * 4 / max(n, 1), 1)  # 4 players active per game
        st.info(
            f"**{n} players · {num_courts} courts · doubles**  \n"
            f"{detail}  \n"
            f"Total: **{total_games} games** · each player plays ~**{avg_games}**"
        )

        st.markdown('<div class="section-label">Player roster</div>', unsafe_allow_html=True)

        # ── Test Fill (dev only) ──────────────────────────────────────────────
        if st.session_state.get("verified_phone") == "7261979719":
            if st.button("🧪 Test Fill", key="btn_test_fill", use_container_width=True):
                _tb = [f"Boy{i+1}"  for i in range(st.session_state.get("num_boys", 6))]
                _tg = [f"Girl{i+1}" for i in range(st.session_state.get("num_girls", 4))]
                st.session_state.boy_names  = _tb
                st.session_state.girl_names = _tg
                for _i, _nm in enumerate(_tb):
                    st.session_state[f"pname_boy_{_i}"]  = _nm
                    st.session_state[f"skill_boy_{_i}"]  = "intermediate"
                for _i, _nm in enumerate(_tg):
                    st.session_state[f"pname_girl_{_i}"] = _nm
                    st.session_state[f"skill_girl_{_i}"] = "intermediate"
                st.session_state.special_instructions = (
                    "Avoid a team with 2 girls if the opponent team has only boys."
                )
                st.rerun()

        # ── Boys ──────────────────────────────────────────────────────────────
        st.markdown('<p class="roster-heading-boys">👦 Boys</p>', unsafe_allow_html=True)
        num_boys = st.number_input(
            "Number of boys", min_value=0, max_value=16, step=1,
            value=st.session_state.get("num_boys", 6),
            key="num_boys_input",
        )
        if num_boys != st.session_state.get("num_boys", 6):
            for _c in range(1, st.session_state.get("num_courts", 2) + 1):
                _hk = f"court_hours_{_c}"
                if _hk in st.session_state:
                    st.session_state[f"_bk_{_hk}"] = st.session_state[_hk]
            cur = st.session_state.get("boy_names", [])
            if num_boys > len(cur):
                cur = cur + [""] * (num_boys - len(cur))
            else:
                cur = cur[:num_boys]
            st.session_state.boy_names = cur
            st.session_state.num_boys  = num_boys
            st.rerun()

        for i in range(st.session_state.get("num_boys", 6)):
            if f"skill_boy_{i}" not in st.session_state:
                st.session_state[f"skill_boy_{i}"] = "intermediate"
            default = (
                st.session_state.get("boy_names", [""] * 20)[i]
                if i < len(st.session_state.get("boy_names", []))
                else ""
            )
            if st.session_state.skill_visible:
                c_name, c_skill = st.columns([5, 4])
                with c_name:
                    st.text_input(
                        f"B{i+1}", value=default,
                        key=f"pname_boy_{i}", placeholder=f"Boy {i+1}",
                        label_visibility="collapsed",
                    )
                with c_skill:
                    st.radio(
                        "Level", options=["intermediate", "beginner"],
                        key=f"skill_boy_{i}", horizontal=True,
                        label_visibility="collapsed",
                    )
            else:
                st.text_input(
                    f"B{i+1}", value=default,
                    key=f"pname_boy_{i}", placeholder=f"Boy {i+1}",
                    label_visibility="collapsed",
                )

        # ── Girls ─────────────────────────────────────────────────────────────
        st.markdown('<p class="roster-heading-girls">👧 Girls</p>', unsafe_allow_html=True)
        num_girls = st.number_input(
            "Number of girls", min_value=0, max_value=8, step=1,
            value=st.session_state.get("num_girls", 4),
            key="num_girls_input",
        )
        if num_girls != st.session_state.get("num_girls", 4):
            for _c in range(1, st.session_state.get("num_courts", 2) + 1):
                _hk = f"court_hours_{_c}"
                if _hk in st.session_state:
                    st.session_state[f"_bk_{_hk}"] = st.session_state[_hk]
            cur = st.session_state.get("girl_names", [])
            if num_girls > len(cur):
                cur = cur + [""] * (num_girls - len(cur))
            else:
                cur = cur[:num_girls]
            st.session_state.girl_names = cur
            st.session_state.num_girls  = num_girls
            st.rerun()

        for i in range(st.session_state.get("num_girls", 4)):
            if f"skill_girl_{i}" not in st.session_state:
                st.session_state[f"skill_girl_{i}"] = "intermediate"
            default = (
                st.session_state.get("girl_names", [""] * 10)[i]
                if i < len(st.session_state.get("girl_names", []))
                else ""
            )
            if st.session_state.skill_visible:
                c_name, c_skill = st.columns([5, 4])
                with c_name:
                    st.text_input(
                        f"G{i+1}", value=default,
                        key=f"pname_girl_{i}", placeholder=f"Girl {i+1}",
                        label_visibility="collapsed",
                    )
                with c_skill:
                    st.radio(
                        "Level", options=["intermediate", "beginner"],
                        key=f"skill_girl_{i}", horizontal=True,
                        label_visibility="collapsed",
                    )
            else:
                st.text_input(
                    f"G{i+1}", value=default,
                    key=f"pname_girl_{i}", placeholder=f"Girl {i+1}",
                    label_visibility="collapsed",
                )

        # ── Special instructions ──────────────────────────────────────────────
        st.markdown('<div class="section-label">Special instructions for schedule generation</div>', unsafe_allow_html=True)
        special_instructions = st.text_area(
            "special_instructions",
            value=st.session_state.get("special_instructions", ""),
            placeholder="e.g. Player 3 and Player 7 should not be on the same team. Keep beginners with at least one intermediate player.",
            height=90,
            label_visibility="collapsed",
        )
        if special_instructions != st.session_state.get("special_instructions", ""):
            st.session_state.special_instructions = special_instructions

        # ── Discreet lock toggle (below player list) ──────────────────────────
        st.markdown(
            """
            <style>
            .skill-toggle-anchor + div,
            .skill-toggle-anchor + div > div,
            .skill-toggle-anchor + div > div > div {
                background: transparent !important;
                border: none !important;
                box-shadow: none !important;
                outline: none !important;
            }
            .skill-toggle-anchor + div button,
            .skill-toggle-anchor + div button:hover,
            .skill-toggle-anchor + div button:focus,
            .skill-toggle-anchor + div button:active {
                min-height: 22px !important;
                height: 22px !important;
                width: auto !important;
                padding: 0 0.4rem !important;
                font-size: 0.6rem !important;
                background: transparent !important;
                border: none !important;
                outline: none !important;
                box-shadow: none !important;
                color: #2e2e3e !important;
                letter-spacing: 0.2em;
            }
            .skill-toggle-anchor + div button:hover {
                color: #555 !important;
            }
            </style>
            <div class="skill-toggle-anchor"></div>
            """,
            unsafe_allow_html=True,
        )
        _vis = st.session_state.skill_visible
        if st.button("· · ·", key="btn_skill_vis", use_container_width=False):
            if _vis:
                st.session_state.skill_visible = False
                st.session_state.show_skill_pw = False
            else:
                st.session_state.show_skill_pw = not st.session_state.show_skill_pw
            st.rerun()

        # ── Password prompt ───────────────────────────────────────────────────
        if st.session_state.show_skill_pw and not st.session_state.skill_visible:
            pw_col, go_col = st.columns([5, 2])
            with pw_col:
                pw_val = st.text_input(
                    "pw", type="password",
                    placeholder="Enter password…",
                    label_visibility="collapsed",
                    key="skill_pw_field",
                )
            with go_col:
                if st.button("Unlock", key="btn_skill_unlock",
                             type="primary", use_container_width=True):
                    if pw_val == _admin_password():
                        st.session_state.skill_visible = True
                        st.session_state.show_skill_pw = False
                        if "skill_pw_field" in st.session_state:
                            del st.session_state["skill_pw_field"]
                    else:
                        st.error("Incorrect password.")
                    st.rerun()

    # ── Generate Schedule tab ───────────────────────────────────────────────
    with tab_s:
        # Read values set in Players tab
        num_courts     = st.session_state.get("num_courts", 2)
        games_per_hour = st.session_state.get("games_per_hour", 5)
        # Rebuild per-court game counts from stored hour sliders
        _ngpc: Dict[int, int] = {}
        for _c in range(1, num_courts + 1):
            _hrs = float(st.session_state.get(f"court_hours_{_c}", 2.0))
            _ngpc[_c] = max(1, round(games_per_hour * _hrs))
        num_games = max(_ngpc.values())  # generate enough rounds for the busiest court

        has_key = bool(_gemini_key())
        use_agent = st.checkbox(
            "Use AI agent (Gemini Flash)",
            value=has_key,
            disabled=not has_key,
            help="Set gemini.api_key in config.yaml to enable." if not has_key
                 else "Uses Gemini Flash ReAct agent for scheduling.",
        )

        st.markdown(" ")

        # ── Generate Schedule ─────────────────────────────────────────────────
        if st.button("🎲 Generate Schedule", type="primary", use_container_width=True):
            st.session_state.show_gen_pw   = True
            st.session_state.show_reset_pw = False
            st.session_state._active_setup_tab = 1
            st.rerun()

        if st.session_state.show_gen_pw:
            st.caption("Enter password to generate schedule")
            pw_col, go_col = st.columns([5, 2])
            with pw_col:
                gen_pw = st.text_input(
                    "gen_pw", type="password", placeholder="Password…",
                    label_visibility="collapsed", key="gen_pw_field",
                )
            with go_col:
                if st.button("Confirm", key="btn_gen_confirm",
                             type="primary", use_container_width=True):
                    if gen_pw == _admin_password():
                        st.session_state.show_gen_pw = False
                        # ── Wipe old schedule/score widget state ─────────────
                        old_schedule = _get().get("schedule", [])
                        for _g in old_schedule:
                            for _prefix in ("sa_", "sb_", "btn_"):
                                st.session_state.pop(f"{_prefix}{_g['game_id']}", None)
                        # ── Collect names & skills from boys + girls sections ─
                        _nb = st.session_state.get("num_boys", 6)
                        _ng = st.session_state.get("num_girls", 4)
                        raw_names = (
                            [(st.session_state.get(f"pname_boy_{i}") or f"Boy {i+1}").strip()  for i in range(_nb)] +
                            [(st.session_state.get(f"pname_girl_{i}") or f"Girl {i+1}").strip() for i in range(_ng)]
                        )
                        raw_skills = (
                            [st.session_state.get(f"skill_boy_{i}", "intermediate")  for i in range(_nb)] +
                            [st.session_state.get(f"skill_girl_{i}", "intermediate") for i in range(_ng)]
                        )
                        seen: Dict[str, int] = {}
                        players: List[str] = []
                        skill_levels: Dict[str, str] = {}
                        for nm, sk in zip(raw_names, raw_skills):
                            nm = nm or "Player"
                            if nm in seen:
                                seen[nm] += 1
                                nm = f"{nm} ({seen[nm]})"
                            else:
                                seen[nm] = 1
                            players.append(nm)
                            skill_levels[nm] = sk
                        # Resolved girl names (after dedup) = last _ng entries
                        resolved_girl_names = players[_nb:]
                        st.session_state.player_names = players
                        st.session_state.girl_names   = resolved_girl_names

                        # ── Detect refine vs fresh ──────────────────────────
                        _saved         = _get()
                        _saved_players = _saved.get("players", [])
                        _saved_courts  = _saved.get("num_courts", 0)
                        _saved_sched   = _saved.get("schedule", [])
                        _is_refine = (
                            bool(_saved_sched)
                            and sorted(_saved_players) == sorted(players)
                            and _saved_courts == num_courts
                        )

                        st.caption("⏱ Hang On there! Might take around 5 mins to generate.")
                        with st.spinner("Generating schedule… 🤖"):
                            try:
                                if use_agent and has_key:
                                    from agent.react_agent import GamePlannerAgent
                                    raw_schedule = GamePlannerAgent().generate_schedule(
                                        players, skill_levels,
                                        num_rounds=num_games, num_courts=num_courts,
                                        special_instructions=st.session_state.get("special_instructions", ""),
                                        previous_schedule=_saved_sched if _is_refine else None,
                                    )
                                    method = "AI agent — refine" if _is_refine else "AI agent — fresh"
                                else:
                                    from services.schedule_service import ScheduleService
                                    raw_schedule = ScheduleService().generate_schedule(
                                        players, skill_levels,
                                        num_rounds=num_games, num_courts=num_courts,
                                    )
                                    method = "algorithm"
                            except Exception as exc:
                                logging.error("Schedule generation failed: %s", exc)
                                from services.schedule_service import ScheduleService
                                raw_schedule = ScheduleService().generate_schedule(
                                    players, skill_levels,
                                    num_rounds=num_games, num_courts=num_courts,
                                )
                                method = "algorithm (fallback)"

                            court_seen: Dict[int, int] = {c: 0 for c in range(1, num_courts + 1)}
                            schedule = []
                            for g in raw_schedule:
                                c = g["court"]
                                limit = _ngpc.get(c, num_games)
                                if court_seen[c] < limit:
                                    schedule.append(g)
                                    court_seen[c] += 1

                            _put({
                                "players":        players,
                                "girl_names":     resolved_girl_names,
                                "skill_levels":   skill_levels,
                                "num_courts":     num_courts,
                                "games_per_hour": games_per_hour,
                                "court_hours": {
                                    c: float(st.session_state.get(f"court_hours_{c}", 2.0))
                                    for c in range(1, num_courts + 1)
                                },
                                "schedule": schedule,
                                "scores": {
                                    g["game_id"]: {"score_a": None, "score_b": None, "submitted": False}
                                    for g in schedule
                                },
                                "special_instructions": st.session_state.get("special_instructions", ""),
                                "session_active": True,
                            })

                        label = "refined" if _is_refine else "generated"
                        st.success(f"✅ {len(schedule)} games {label} via {method}")
                        st.session_state._active_setup_tab = 1
                        st.rerun()
                    else:
                        st.error("Incorrect password.")

        # ── Reset Session ─────────────────────────────────────────────────────
        st.markdown(" ")
        if st.button("🔄 Reset Session", use_container_width=True):
            st.session_state.show_reset_pw = True
            st.session_state.show_gen_pw   = False
            st.session_state._active_setup_tab = 1
            st.rerun()

        if st.session_state.show_reset_pw:
            st.caption("Enter password to reset session")
            rp_col, rgo_col = st.columns([5, 2])
            with rp_col:
                reset_pw = st.text_input(
                    "reset_pw", type="password", placeholder="Password…",
                    label_visibility="collapsed", key="reset_pw_field",
                )
            with rgo_col:
                if st.button("Confirm", key="btn_reset_confirm",
                             type="primary", use_container_width=True):
                    if reset_pw == _admin_password():
                        st.session_state.show_reset_pw = False
                        # Clear all per-game widget state (scores, winner selections)
                        for k in list(st.session_state.keys()):
                            if k.startswith(("sa_", "sb_", "winner_", "win_a_", "win_b_", "court_hours_")):
                                del st.session_state[k]
                        # Reset setup widgets to defaults
                        st.session_state.num_boys      = 6
                        st.session_state.num_girls     = 4
                        st.session_state.num_players   = 10
                        st.session_state.num_courts    = 2
                        st.session_state.games_per_hour = 5
                        st.session_state.boy_names     = [""] * 6
                        st.session_state.girl_names    = [""] * 4
                        st.session_state.player_names  = [""] * 10
                        st.session_state.special_instructions = ""
                        for _i in range(20):
                            st.session_state.pop(f"pname_{_i}", None)
                            st.session_state.pop(f"pname_boy_{_i}", None)
                            st.session_state.pop(f"pname_girl_{_i}", None)
                        with st.spinner("Resetting…"):
                            _put(_empty_state())
                        st.info("Tournament reset. Please refresh your browser to start fresh.")
                        st.session_state._active_setup_tab = 1
                        st.rerun()
                    else:
                        st.error("Incorrect password.")

        # ── Schedule preview — always read fresh state so submitted scores show ──
        state = _get()
        if not state.get("schedule"):
            st.info("No schedule yet — click **Generate Schedule** above.")
        else:
            schedule = state["schedule"]
            scores   = state.get("scores", {})

            st.divider()
            st.subheader(f"📋 Schedule  ·  {len(schedule)} games")

            dl_col1, dl_col2 = st.columns(2)
            with dl_col1:
                st.download_button(
                    "⬇️ Download Schedule (Word)",
                    data=_build_docx(schedule, scores),
                    file_name="game_schedule.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with dl_col2:
                st.download_button(
                    "⬇️ Download Schedule (Excel)",
                    data=_build_xlsx(schedule, scores),
                    file_name="game_schedule.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            dl_col3, _ = st.columns(2)
            with dl_col3:
                st.download_button(
                    "⬇️ Download Schedule (CSV)",
                    data=_build_csv(schedule, scores),
                    file_name="game_schedule.csv",
                    mime="text/csv",
                    use_container_width=True,
                )

            # ── Upload revised schedule ───────────────────────────────────────
            with st.expander("📤 Upload Revised Schedule (Excel)", expanded=False):
                st.caption(
                    "Download the Excel above, edit player names / courts / scores offline, "
                    "then upload it here to replace the current schedule."
                )
                uploaded_xl = st.file_uploader(
                    "Choose Excel file",
                    type=["xlsx"],
                    key="schedule_upload",
                    label_visibility="collapsed",
                )
                if uploaded_xl is not None:
                    parsed_sched, parsed_scores, parse_err = _parse_uploaded_xlsx(uploaded_xl.read())
                    if parse_err:
                        st.error(parse_err)
                    else:
                        submitted_count = sum(
                            1 for v in parsed_scores.values() if v.get("submitted")
                        )
                        st.success(
                            f"✓ {len(parsed_sched)} games parsed · "
                            f"{submitted_count} with scores · "
                            f"{len(parsed_sched) - submitted_count} pending"
                        )
                        _render_table(parsed_sched, parsed_scores, show_court=True)
                        st.caption("Enter admin password to apply this schedule")
                        up_pw_col, up_go_col = st.columns([5, 2])
                        with up_pw_col:
                            up_pw = st.text_input(
                                "up_pw", type="password", placeholder="Password…",
                                label_visibility="collapsed", key="upload_pw_field",
                            )
                        with up_go_col:
                            if st.button(
                                "✅ Apply", key="btn_upload_apply",
                                type="primary", use_container_width=True,
                            ):
                                if up_pw == _admin_password():
                                    new_state = copy.deepcopy(_get())
                                    new_state["schedule"]       = parsed_sched
                                    new_state["scores"]         = parsed_scores
                                    new_state["critics_choice"] = None
                                    new_state["num_courts"]     = max(
                                        (g["court"] for g in parsed_sched), default=2
                                    )
                                    with st.spinner("Saving revised schedule…"):
                                        _put(new_state)
                                    st.success("Schedule updated. Refresh the court pages.")
                                    st.rerun()
                                else:
                                    st.error("Incorrect password.")

            num_courts_now = state.get("num_courts", 2)
            sched_tab_labels = [f"🏟 Court {c}" for c in range(1, num_courts_now + 1)] + ["📄 All"]
            sched_tabs = st.tabs(sched_tab_labels)
            for c, tab in enumerate(sched_tabs[:-1], start=1):
                with tab:
                    _render_table([g for g in schedule if g["court"] == c], scores)
            with sched_tabs[-1]:
                _render_table(schedule, scores, show_court=True)



def _render_table(games: List[dict], scores: dict, show_court: bool = False) -> None:
    if not games:
        st.info("No games.")
        return
    rows = []
    for game_num, g in enumerate(games, start=1):
        sd   = scores.get(g["game_id"], {})
        done = sd.get("submitted", False)
        row = {
            "":      "✅" if done else "⏳",
            "Game":  game_num,
        }
        if show_court:
            row["Court"] = g.get("court", "")
        row["Team A"]    = " & ".join(g["team_a"])
        row["Score A"]   = str(sd["score_a"]) if done else "—"
        row["Team B"]    = " & ".join(g["team_b"])
        row["Score B"]   = str(sd["score_b"]) if done else "—"
        row["Rest"]      = ", ".join(g.get("sitting_out", []))
        rows.append(row)
    st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)


def _build_docx(schedule: List[dict], scores: dict = None) -> bytes:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    scores = scores or {}
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)

    title = doc.add_heading("Game Schedule", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 6 columns: Game | Court | Team A | Score A | Team B | Score B
    headers = ["Game", "Court", "Team A", "Score A", "Team B", "Score B"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # One game = 2 rows (one per player), Game/Court/ScoreA/ScoreB merged vertically
    for game_num, game in enumerate(schedule, start=1):
        ta   = game.get("team_a", [])
        tb   = game.get("team_b", [])
        sd   = scores.get(game.get("game_id", ""), {})
        done = sd.get("submitted", False)
        score_a = str(sd["score_a"]) if done else ""
        score_b = str(sd["score_b"]) if done else ""

        r1 = table.add_row().cells
        r1[0].text = f"Game {game_num}"
        r1[1].text = f"Court {game.get('court', '')}"
        r1[2].text = ta[0] if ta else ""
        r1[3].text = score_a
        r1[4].text = tb[0] if tb else ""
        r1[5].text = score_b

        r2 = table.add_row().cells
        r2[2].text = ta[1] if len(ta) > 1 else ""
        r2[4].text = tb[1] if len(tb) > 1 else ""

        # Merge Game, Court, ScoreA, ScoreB across 2 rows
        for col_idx in (0, 1, 3, 5):
            r1[col_idx].merge(r2[col_idx])
            r1[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in (r1, r2):
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

    col_widths = [Cm(2.0), Cm(2.0), Cm(4.5), Cm(1.8), Cm(4.5), Cm(1.8)]
    for row in table.rows:
        for cell, w in zip(row.cells, col_widths):
            cell.width = w

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _parse_uploaded_xlsx(file_bytes: bytes):
    """Parse an uploaded schedule Excel into (schedule, scores, error_string).

    Expects the same columns produced by _build_xlsx:
      Game, Court, Team A, Score A, Team B, Score B, Sitting Out
    Returns (None, None, error_msg) on failure.
    """
    import io
    import math
    import pandas as pd

    def _safe_str(val) -> str:
        if val is None:
            return ""
        try:
            if math.isnan(float(val)):
                return ""
        except (TypeError, ValueError):
            pass
        return str(val).strip()

    try:
        try:
            df = pd.read_excel(io.BytesIO(file_bytes), sheet_name="Schedule")
        except Exception:
            df = pd.read_excel(io.BytesIO(file_bytes))
    except Exception as exc:
        return None, None, f"Could not read file: {exc}"

    required_cols = {"Court", "Team A", "Team B"}
    missing = required_cols - set(df.columns)
    if missing:
        return None, None, f"Missing required columns: {', '.join(sorted(missing))}"

    schedule: List[dict] = []
    scores: dict = {}
    seen_ids: set = set()
    errors: List[str] = []

    round_num = 0
    prev_court = 0

    for row_num, (_, row) in enumerate(df.iterrows()):
        court_raw = _safe_str(row.get("Court", ""))
        if not court_raw:
            continue
        try:
            court = int(float(court_raw))
        except ValueError:
            errors.append(f"Row {row_num + 2}: invalid Court value '{court_raw}'")
            continue

        # Detect new round whenever court resets to a lower/equal value
        if row_num == 0 or court <= prev_court:
            round_num += 1
        prev_court = court

        team_a_str = _safe_str(row.get("Team A", ""))
        team_b_str = _safe_str(row.get("Team B", ""))
        sitting_str = _safe_str(row.get("Sitting Out", ""))

        team_a = [n.strip() for n in team_a_str.split("&") if n.strip()]
        team_b = [n.strip() for n in team_b_str.split("&") if n.strip()]
        sitting = [n.strip() for n in sitting_str.split(",") if n.strip()]

        if len(team_a) != 2 or len(team_b) != 2:
            errors.append(
                f"Row {row_num + 2}: Team A or Team B must have exactly 2 players "
                f"(use 'Player1 & Player2' format). Got: '{team_a_str}' vs '{team_b_str}'"
            )
            continue

        # Ensure unique game_id
        game_id = f"r{round_num}_c{court}"
        suffix = 1
        while game_id in seen_ids:
            game_id = f"r{round_num}_c{court}_{suffix}"
            suffix += 1
        seen_ids.add(game_id)

        schedule.append({
            "round":       round_num,
            "court":       court,
            "team_a":      team_a,
            "team_b":      team_b,
            "sitting_out": sitting,
            "time_slot":   f"{(round_num - 1) * 10}–{round_num * 10} min",
            "game_id":     game_id,
        })

        score_a_raw = _safe_str(row.get("Score A", ""))
        score_b_raw = _safe_str(row.get("Score B", ""))
        if score_a_raw and score_b_raw:
            try:
                scores[game_id] = {
                    "score_a":   int(float(score_a_raw)),
                    "score_b":   int(float(score_b_raw)),
                    "submitted": True,
                }
            except ValueError:
                pass  # malformed scores → treat as not submitted

    if errors:
        return None, None, "\n".join(errors)
    if not schedule:
        return None, None, "No valid game rows found in the file."

    return schedule, scores, None


def _build_xlsx(schedule: List[dict], scores: dict = None) -> bytes:
    import io
    import pandas as pd

    scores = scores or {}
    rows = []
    for game_num, g in enumerate(schedule, start=1):
        sd   = scores.get(g.get("game_id", ""), {})
        done = sd.get("submitted", False)
        rows.append({
            "Game":        game_num,
            "Court":       g.get("court", ""),
            "Team A":      " & ".join(g.get("team_a", [])),
            "Score A":     sd["score_a"] if done else "",
            "Team B":      " & ".join(g.get("team_b", [])),
            "Score B":     sd["score_b"] if done else "",
            "Sitting Out": ", ".join(g.get("sitting_out", [])),
        })

    df = pd.DataFrame(rows)

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Schedule")

        ws = writer.sheets["Schedule"]

        # Column widths: A=Game, B=Court, C=TeamA, D=ScoreA, E=TeamB, F=ScoreB, G=SittingOut
        col_widths = {"A": 8, "B": 8, "C": 26, "D": 10, "E": 26, "F": 10, "G": 28}
        for col, width in col_widths.items():
            ws.column_dimensions[col].width = width

        from openpyxl.styles import Font, PatternFill, Alignment
        header_fill = PatternFill("solid", fgColor="1F4E79")
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        # Centre-align: Game(1), Court(2), ScoreA(4), ScoreB(6)
        centre_cols = {1, 2, 4, 6}
        light = PatternFill("solid", fgColor="EBF3FB")
        for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
            fill = light if row_idx % 2 == 0 else PatternFill()
            for col_idx, cell in enumerate(row, start=1):
                cell.fill = fill
                if col_idx in centre_cols:
                    cell.alignment = Alignment(horizontal="center")

    buf.seek(0)
    return buf.getvalue()


def _build_csv(schedule: List[dict], scores: dict = None) -> bytes:
    import io
    import pandas as pd

    scores = scores or {}
    rows = []
    for game_num, g in enumerate(schedule, start=1):
        sd   = scores.get(g.get("game_id", ""), {})
        done = sd.get("submitted", False)
        rows.append({
            "Game":        game_num,
            "Court":       g.get("court", ""),
            "Team A":      " & ".join(g.get("team_a", [])),
            "Score A":     sd["score_a"] if done else "",
            "Team B":      " & ".join(g.get("team_b", [])),
            "Score B":     sd["score_b"] if done else "",
            "Sitting Out": ", ".join(g.get("sitting_out", [])),
        })

    buf = io.BytesIO()
    pd.DataFrame(rows).to_csv(buf, index=False)
    buf.seek(0)
    return buf.getvalue()


# ===========================================================================
# Page: Court  (mobile-first score entry)
# ===========================================================================

def show_court(court: int) -> None:
    st.markdown(
        f'<div class="page-header">'
        f'<h1>🏟 Court {court}</h1>'
        f'<p>Enter scores as each game finishes</p>'
        f'</div>',
        unsafe_allow_html=True,
    )

    state = _get()
    if not state.get("schedule"):
        st.warning("No schedule yet — go to **Setup** first.")
        return

    court_games = [g for g in state["schedule"] if g["court"] == court]
    scores      = state["scores"]

    if not court_games:
        st.info(f"No games for Court {court}.")
        return

    done  = sum(1 for g in court_games if scores.get(g["game_id"], {}).get("submitted"))
    total = len(court_games)

    c1, c2, c3 = st.columns(3)
    c1.metric("Total",  total)
    c2.metric("Done",   done)
    c3.metric("Left",   total - done)
    st.progress(done / total if total else 0)
    st.divider()

    for game_num, game in enumerate(court_games, start=1):
        gid = game["game_id"]
        sd  = scores.get(gid, {})
        submitted = sd.get("submitted", False)

        # ── Pre-initialise session state for scores & winner ──────────────
        if f"sa_{gid}" not in st.session_state:
            st.session_state[f"sa_{gid}"] = int(sd["score_a"]) if submitted and sd.get("score_a") is not None else 0
        if f"sb_{gid}" not in st.session_state:
            st.session_state[f"sb_{gid}"] = int(sd["score_b"]) if submitted and sd.get("score_b") is not None else 0
        if f"winner_{gid}" not in st.session_state:
            if submitted:
                _sa, _sb = sd.get("score_a") or 0, sd.get("score_b") or 0
                st.session_state[f"winner_{gid}"] = (
                    "Team A" if _sa > _sb else "Team B" if _sb > _sa else "—"
                )
            else:
                st.session_state[f"winner_{gid}"] = "—"

        winner = st.session_state.get(f"winner_{gid}", "—")

        icon = "✅" if submitted else "⏳"
        with st.expander(f"{icon}  Game {game_num}", expanded=not submitted):
            # Team A row: card + Win button
            col_card_a, col_btn_a = st.columns([5, 2])
            with col_card_a:
                st.markdown(
                    f'<div class="team-card-a">'
                    f'<strong>Team A</strong> &nbsp; {" &amp; ".join(game["team_a"])}'
                    f'</div>',
                    unsafe_allow_html=True,
                )
            with col_btn_a:
                if st.button(
                    "🏆 Win" if winner != "Team A" else "✅ Won",
                    key=f"win_a_{gid}",
                    type="primary" if winner == "Team A" else "secondary",
                    use_container_width=True,
                ):
                    st.session_state[f"winner_{gid}"] = "Team A"
                    st.session_state[f"sa_{gid}"] = max(11, st.session_state.get(f"sa_{gid}", 0))
                    st.rerun()

            # Team B row: card + Win button
            col_card_b, col_btn_b = st.columns([5, 2])
            with col_card_b:
                st.markdown(
                    f'<div class="team-card-b">'
                    f'<strong>Team B</strong> &nbsp; {" &amp; ".join(game["team_b"])}'
                    f'</div>',
                    unsafe_allow_html=True,
                )
            with col_btn_b:
                if st.button(
                    "🏆 Win" if winner != "Team B" else "✅ Won",
                    key=f"win_b_{gid}",
                    type="primary" if winner == "Team B" else "secondary",
                    use_container_width=True,
                ):
                    st.session_state[f"winner_{gid}"] = "Team B"
                    st.session_state[f"sb_{gid}"] = max(11, st.session_state.get(f"sb_{gid}", 0))
                    st.rerun()

            # Score inputs — winner defaults to 11, max 30 for tie-breaks
            col_a, col_b = st.columns(2)
            with col_a:
                score_a = st.number_input(
                    "Team A score", min_value=0, max_value=30,
                    key=f"sa_{gid}",
                )
            with col_b:
                score_b = st.number_input(
                    "Team B score", min_value=0, max_value=30,
                    key=f"sb_{gid}",
                )

            btn = "✏️ Update Score" if submitted else "✅ Submit Score"
            if st.button(btn, key=f"btn_{gid}", type="primary", use_container_width=True):
                new_state = copy.deepcopy(_get())
                new_state["scores"][gid] = {
                    "score_a": score_a, "score_b": score_b, "submitted": True
                }
                with st.spinner("Saving…"):
                    _put(new_state)
                st.toast("Score saved!")
                st.rerun()

            # Result banner
            if submitted:
                sa, sb = sd["score_a"], sd["score_b"]
                if sa > sb:
                    st.success(f"🏆 Team A wins!  {sa} – {sb}")
                elif sb > sa:
                    st.success(f"🏆 Team B wins!  {sb} – {sa}")
                else:
                    st.success(f"🤝 Draw!  {sa} – {sb}")




# ===========================================================================
# Page: Leaderboard
# ===========================================================================

def _get_critics_choice(
    lb: list,
    special_instructions: str,  # kept for call-site compat, not used in prompt
    podium_names: list = None,
    girl_names: set = None,
) -> Optional[list]:
    """Call Gemini to pick Critic's Choice top 3.

    Rules enforced both in the prompt and in post-processing:
    - At most 1 player may overlap with the Winner Podium (podium_names).
    - At least 1 girl must be included (when girls are available).
    - The girl chosen must NOT be the same girl already on the Winner Podium.
    """
    import json as _json
    import google.generativeai as genai
    from agent.react_agent import _load_gemini_config, _load_api_key

    gcfg    = _load_gemini_config()
    api_key = _load_api_key()
    if not api_key:
        raise ValueError("Gemini API key not found — check secrets.toml or config.yaml.")

    podium_names  = list(podium_names or [])
    girl_names    = set(girl_names or [])
    podium_set    = set(podium_names)
    winner_girls  = girl_names & podium_set          # girls already on winner podium
    available_girls = girl_names - podium_set        # girls NOT on winner podium

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(
        model_name=gcfg.get("model_name", "gemini-2.0-flash"),
        generation_config=genai.GenerationConfig(
            temperature=float(gcfg.get("temperature", 0.7)),
            top_p=float(gcfg.get("top_p", 0.95)),
            response_mime_type="application/json",
        ),
    )

    stats_lines = "\n".join(
        f"- {p['name']}: {p['games_won']}W / {p['games_lost']}L, "
        f"For: {p['points_gained']}, Against: {p['points_conceded']}, "
        f"Net: {p['net_points']:+d}, Win rate: {p['win_rate']}%"
        for p in lb
    )

    # Build constraint block for the prompt
    constraint_lines = []
    if podium_names:
        constraint_lines.append(
            f"WINNER PODIUM (by avg points) players are: {', '.join(podium_names)}. "
            "AT MOST 1 of these players may appear in your Critic's Choice podium. "
            "The other 2 picks MUST come from outside this list."
        )
    if girl_names:
        if available_girls:
            constraint_lines.append(
                f"Girl players NOT on the winner podium: {', '.join(sorted(available_girls))}. "
                "You MUST include at least 1 of these girls in your top 3. "
                + (
                    f"Do NOT pick {', '.join(sorted(winner_girls))} as your girl pick "
                    f"— {'she is' if len(winner_girls) == 1 else 'they are'} already on the winner podium."
                    if winner_girls else ""
                )
            )
        elif girl_names:
            constraint_lines.append(
                f"Girl players: {', '.join(sorted(girl_names))}. "
                "Include at least 1 girl in your top 3."
            )

    constraints_block = (
        "\n\nMANDATORY CONSTRAINTS (non-negotiable):\n" + "\n".join(f"{i+1}. {c}" for i, c in enumerate(constraint_lines))
        if constraint_lines else ""
    )

    prompt = (
        "You are a sharp, fair game critic reviewing a doubles sports tournament.\n\n"
        "PLAYER STATISTICS:\n"
        + stats_lines
        + constraints_block
        + "\n\nYOUR TASK — select the Critic's (AI) Choice top 3 podium.\n\n"
        "SELECTION CRITERIA:\n"
        "1. Look beyond raw wins — reward bravery, consistency, close scorelines, and standout rallies.\n"
        "2. A girl who played hard against stronger opponents deserves recognition — "
        "high points_gained in losses shows courage.\n\n"
        "Write a 2-sentence reason (max 35 words) for each pick. "
        "CRITICAL: Only mention POSITIVE achievements — great points scored, impressive wins, "
        "strong rallies, standout performances. "
        "Do NOT mention losses, loss counts, win/loss percentages, gender, or any negative stats. "
        "The tone must be celebratory and encouraging throughout.\n\n"
        'Respond ONLY in valid JSON, no markdown, no code fences, no extra text:\n'
        '{"podium": [{"rank": 1, "name": "ExactName", "reason": "..."}, '
        '{"rank": 2, "name": "ExactName", "reason": "..."}, '
        '{"rank": 3, "name": "ExactName", "reason": "..."}]}\n\n'
        "Names must exactly match names from the stats list."
    )

    response = model.generate_content(prompt)
    text = response.text.strip()
    if "```" in text:
        parts = text.split("```")
        text = parts[1] if len(parts) > 1 else parts[0]
        if text.startswith("json"):
            text = text[4:]
    text = text.strip()

    data        = _json.loads(text)
    podium      = data.get("podium", [])
    valid_names = {p["name"] for p in lb}
    podium      = [e for e in podium if e.get("name") in valid_names][:3]

    if len(podium) < 3:
        raise ValueError(f"Gemini returned fewer than 3 valid picks: {podium}")

    # ── Post-processing validation ────────────────────────────────────────────
    # Rule 1: at most 1 overlap with winner podium
    overlap = [e["name"] for e in podium if e["name"] in podium_set]
    if len(overlap) > 1:
        raise ValueError(
            f"Critic's Choice overlaps winner podium on {overlap} — max 1 allowed. "
            "Please try again."
        )

    # Rule 2 & 3: at least 1 girl, and not the same girl as winner podium
    if girl_names:
        cc_girls = [e["name"] for e in podium if e["name"] in girl_names]
        if not cc_girls:
            raise ValueError("Critic's Choice has no girls — at least 1 required. Please try again.")
        if available_girls:
            bad_girls = [n for n in cc_girls if n in winner_girls]
            if bad_girls:
                raise ValueError(
                    f"{', '.join(bad_girls)} is already on the winner podium — "
                    "the girl pick must be different. Please try again."
                )

    return podium


def _identify_girls(player_names: list, special_instructions: str) -> set:
    """Return player names explicitly introduced as girls in special_instructions.

    Only names that appear AFTER a gender keyword within the same sentence/clause
    are included, preventing false positives from names that merely appear nearby.
    """
    import re as _re
    if not special_instructions:
        return set()

    name_map  = {n.lower(): n for n in player_names}
    girls: set = set()
    gender_re  = _re.compile(r'\b(?:girls?|females?|women|woman|ladies|lady)\b')

    # Split into clauses on sentence-ending punctuation or newlines
    clauses = _re.split(r'[.;!\n]+', special_instructions.lower())

    for clause in clauses:
        gm = gender_re.search(clause)
        if not gm:
            continue
        # Only consider names that appear AFTER the gender keyword in this clause
        after = clause[gm.start():]
        for name_lower, name_orig in name_map.items():
            if name_lower in after:
                girls.add(name_orig)

        # Also handle reverse pattern: "[Name1, Name2] are girls"
        before_match = _re.match(
            r'^(.*?)\bare\s+(?:a\s+)?(?:girls?|females?|women|woman|ladies?)\b',
            clause,
        )
        if before_match:
            before_text = before_match.group(1)
            for name_lower, name_orig in name_map.items():
                if name_lower in before_text:
                    girls.add(name_orig)

    return girls


def show_all_time_leaderboard() -> None:
    from datetime import datetime, timedelta, timezone

    st.markdown(
        '<div class="page-header">'
        '<h1>📊 All-Time Leaderboard</h1>'
        '<p>Aggregate standings · last 30 days</p>'
        '</div>',
        unsafe_allow_html=True,
    )

    if st.button("🔄 Refresh", use_container_width=True):
        st.rerun()

    pub = _github_load_file("data/published_results.json", default={"results": []})
    all_results = pub.get("results", []) if isinstance(pub, dict) else []

    cutoff = datetime.now(timezone.utc) - timedelta(days=30)
    cutoff_naive = cutoff.replace(tzinfo=None)
    filtered = []
    for entry in all_results:
        try:
            ts = datetime.strptime(entry["timestamp"], "%Y-%m-%d %H:%M:%S UTC")
            if ts >= cutoff_naive:
                filtered.append(entry)
        except (KeyError, ValueError):
            pass

    if not filtered:
        st.info("No published results in the last 30 days. Publish a session from the Leaderboard page to start building history.")
        return

    # ── Aggregate per-player stats across all filtered sessions ──────────────
    girl_names_all: set = set()
    agg: dict = {}

    for entry in filtered:
        girl_names_all.update(entry.get("girl_names", []))
        for p in entry.get("players", []):
            name = p["name"]
            if name not in agg:
                agg[name] = {"name": name, "wins": 0, "losses": 0, "net_points": 0}
            agg[name]["wins"]       += p.get("wins", 0)
            agg[name]["losses"]     += p.get("losses", 0)
            agg[name]["net_points"] += p.get("net_points", 0)

    for stat in agg.values():
        stat["games_played"] = stat["wins"] + stat["losses"]

    active = [p for p in agg.values() if p["games_played"] > 0]

    if not active:
        st.info("No player data found in the filtered results.")
        return

    # ── Podium helpers (same as daily leaderboard) ────────────────────────────
    def _avg_pts(p: dict) -> float:
        gp = p.get("games_played", 0)
        return (p.get("wins", 0) * 2) / gp if gp > 0 else 0.0

    def _rank_key(p: dict):
        return (_avg_pts(p), p.get("net_points", 0))

    girls_pool = sorted([p for p in active if p["name"] in girl_names_all], key=_rank_key, reverse=True)
    boys_pool  = sorted([p for p in active if p["name"] not in girl_names_all], key=_rank_key, reverse=True)

    if len(girls_pool) >= 2 and len(boys_pool) >= 2:
        selected = boys_pool[:2] + girls_pool[:1]
    else:
        selected = sorted(active, key=_rank_key, reverse=True)[:3]

    podium_players = sorted(selected, key=_rank_key, reverse=True)

    # ── Render podium ─────────────────────────────────────────────────────────
    st.markdown(
        '<div style="background:#F9A825;border-radius:10px;padding:0.75rem 1rem;'
        'text-align:center;font-size:1rem;font-weight:700;color:#1A1200;'
        'margin-bottom:0.75rem;letter-spacing:0.01em;">'
        '🏆 All-Time Winner Podium (by avg points)</div>',
        unsafe_allow_html=True,
    )

    _medal_border = ["#F9A825", "#9E9E9E", "#EF5350"]
    _BOY_BG  = "#0A1628"
    _GIRL_BG = "#1E0818"
    _GIRL_SVG = '<span style="font-size:2.4rem;line-height:1">👧</span>'
    _BOY_SVG  = '<span style="font-size:2.4rem;line-height:1">👱</span>'

    def _medal_svg(fill: str, stroke: str) -> str:
        return (
            '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 40 50" width="38" height="38">'
            f'<rect x="16" y="0" width="8" height="14" fill="{stroke}" rx="3"/>'
            f'<circle cx="20" cy="32" r="17" fill="{fill}" stroke="{stroke}" stroke-width="2"/>'
            f'<circle cx="20" cy="32" r="11" fill="none" stroke="{stroke}" stroke-width="1.5" opacity="0.45"/>'
            '</svg>'
        )

    _MEDAL_SVGS = [
        _medal_svg("#FFD700", "#B8860B"),
        _medal_svg("#D8D8D8", "#888888"),
        _medal_svg("#CD7F32", "#8B4513"),
    ]

    cols = st.columns(len(podium_players))
    for col, medal_svg, p, border in zip(cols, _MEDAL_SVGS, podium_players, _medal_border):
        is_girl = p["name"] in girl_names_all
        icon = _GIRL_SVG if is_girl else _BOY_SVG
        bg   = _GIRL_BG if is_girl else _BOY_BG
        avg  = _avg_pts(p)
        with col:
            st.markdown(
                f'<div style="background:{bg};border-top:4px solid {border};'
                f'border-radius:12px;padding:0.9rem 0.6rem;text-align:center;">'
                f'<div style="line-height:1;margin-bottom:0.1rem">{medal_svg}</div>'
                f'<div style="line-height:1;margin-bottom:0.2rem">{icon}</div>'
                f'<div style="font-weight:700;font-size:0.95rem;margin-top:0.3rem;color:#E8EAF0;">'
                f'{p["name"]}</div>'
                f'<div style="font-size:0.72rem;color:#F9A825;margin-top:0.2rem;font-weight:600;">'
                f'avg {avg:.1f} pts/game</div>'
                f'<div style="font-size:0.65rem;color:#aaa;margin-top:0.3rem;">'
                f'{p["wins"]}W · {p["losses"]}L · net {p["net_points"]:+d}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

    st.divider()

    # ── Full table ────────────────────────────────────────────────────────────
    ranked = sorted(active, key=_rank_key, reverse=True)
    rows = [
        {
            "#":       i + 1,
            "Player":  p["name"],
            "Avg Pts": round(_avg_pts(p), 2),
            "W":       p["wins"],
            "L":       p["losses"],
            "GP":      p["games_played"],
            "Net":     p["net_points"],
        }
        for i, p in enumerate(ranked)
    ]
    st.dataframe(
        pd.DataFrame(rows),
        use_container_width=True,
        hide_index=True,
        column_config={
            "Net":     st.column_config.NumberColumn("Net", format="%+d"),
            "Avg Pts": st.column_config.NumberColumn("Avg Pts", format="%.2f"),
        },
    )

    st.caption(f"Data from {len(filtered)} session(s) published in the last 30 days.")


def show_leaderboard() -> None:
    st.markdown(
        '<div class="page-header">'
        '<h1>🏆 Leaderboard</h1>'
        '<p>Live standings · doubles tournament</p>'
        '</div>',
        unsafe_allow_html=True,
    )

    if st.button("🔄 Refresh", use_container_width=True):
        st.rerun()

    state = _get()

    # ── Load published results history from GitHub ────────────────────────────
    _pub = _github_load_file("data/published_results.json", default={"results": []})
    published_history = (_pub.get("results", []) if isinstance(_pub, dict) else [])
    # published_history is a list of dicts: [{timestamp, players: [{name, wins, losses, net_points}]}]
    # UI representation to be built later

    if not state.get("schedule"):
        st.info("No schedule yet — go to **Setup** first.")
        return

    from services.leaderboard_service import LeaderboardService
    lb    = LeaderboardService().calculate_leaderboard(state["schedule"], state["scores"])
    done  = sum(1 for v in state["scores"].values() if v.get("submitted"))
    total = len(state["scores"])

    if total:
        st.progress(done / total, text=f"{done} / {total} games  ({int(done/total*100)}%)")

    if not lb:
        st.info("No scores yet — enter them on the Court pages.")
        return

    # ── Winner Podium (by avg points per game) ───────────────────────────────
    st.markdown(
        '<div style="background:#F9A825;border-radius:10px;padding:0.75rem 1rem;'
        'text-align:center;font-size:1rem;font-weight:700;color:#1A1200;'
        'margin-bottom:0.75rem;letter-spacing:0.01em;">'
        '🏆 Winner Podium (by avg points)</div>',
        unsafe_allow_html=True,
    )

    def _avg_pts(p: dict) -> float:
        # avg league-points per game: (wins × 2) / games_played
        gp = p.get("games_played", 0)
        return (p.get("games_won", 0) * 2) / gp if gp > 0 else 0.0

    def _rank_key(p: dict):
        return (_avg_pts(p), p.get("net_points", 0))

    active_for_podium = [p for p in lb if p.get("games_played", 0) > 0]
    all_names  = [p["name"] for p in active_for_podium]
    _saved_girls = set(state.get("girl_names", []))
    girl_names   = _saved_girls if _saved_girls else _identify_girls(all_names, state.get("special_instructions", ""))

    girls_pool = sorted([p for p in active_for_podium if p["name"] in girl_names],
                        key=_rank_key, reverse=True)
    boys_pool  = sorted([p for p in active_for_podium if p["name"] not in girl_names],
                        key=_rank_key, reverse=True)

    # 2 boys + 1 girl only when ≥2 girls are in the tournament
    if len(girls_pool) >= 2 and len(boys_pool) >= 2:
        selected = boys_pool[:2] + girls_pool[:1]
    else:
        selected = sorted(active_for_podium, key=_rank_key, reverse=True)[:3]

    podium_players = sorted(selected, key=_rank_key, reverse=True)

    _medal_border = ["#F9A825", "#9E9E9E", "#EF5350"]
    _BOY_BG  = "#0A1628"   # dark navy
    _GIRL_BG = "#1E0818"   # dark rose

    _GIRL_SVG = '<span style="font-size:2.4rem;line-height:1">👧</span>'
    _BOY_SVG  = '<span style="font-size:2.4rem;line-height:1">👱</span>'

    def _medal_svg(fill: str, stroke: str) -> str:
        return (
            '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 40 50" width="38" height="38">'
            f'<rect x="16" y="0" width="8" height="14" fill="{stroke}" rx="3"/>'
            f'<circle cx="20" cy="32" r="17" fill="{fill}" stroke="{stroke}" stroke-width="2"/>'
            f'<circle cx="20" cy="32" r="11" fill="none" stroke="{stroke}" stroke-width="1.5" opacity="0.45"/>'
            '</svg>'
        )

    _MEDAL_SVGS = [
        _medal_svg("#FFD700", "#B8860B"),   # gold
        _medal_svg("#D8D8D8", "#888888"),   # silver
        _medal_svg("#CD7F32", "#8B4513"),   # bronze
    ]

    raw_quips = random.sample(_ALL_QUIPS, min(len(podium_players), len(_ALL_QUIPS)))

    cols = st.columns(len(podium_players))
    for col, medal_svg, p, border, raw_quip in zip(
        cols, _MEDAL_SVGS, podium_players, _medal_border, raw_quips
    ):
        quip    = raw_quip.format(name=p["name"])
        is_girl = p["name"] in girl_names
        icon    = _GIRL_SVG if is_girl else _BOY_SVG
        bg      = _GIRL_BG if is_girl else _BOY_BG
        avg     = _avg_pts(p)
        with col:
            st.markdown(
                f'<div style="background:{bg};border-top:4px solid {border};'
                f'border-radius:12px;padding:0.9rem 0.6rem;text-align:center;">'
                f'<div style="line-height:1;margin-bottom:0.1rem">{medal_svg}</div>'
                f'<div style="line-height:1;margin-bottom:0.2rem">{icon}</div>'
                f'<div style="font-weight:700;font-size:0.95rem;margin-top:0.3rem;color:#E8EAF0;">'
                f'{p["name"]}</div>'
                f'<div style="font-size:0.72rem;color:#F9A825;margin-top:0.2rem;font-weight:600;">'
                f'avg {avg:.1f} pts/game</div>'
                f'<div style="font-size:0.75rem;color:#aaa;margin-top:0.4rem;'
                f'font-style:italic;line-height:1.4">{quip}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

    st.divider()

    # ── Critic's (AI) Choice Podium ───────────────────────────────────────────
    active_lb = [p for p in lb if p.get("games_played", 0) > 0]
    special_instructions = state.get("special_instructions", "")

    if st.button("🎭 Critic's Choice Podium  ᴬᴵ ᵍᵉⁿᵉʳᵃᵗᵉᵈ", type="primary", use_container_width=True):
        if len(active_lb) >= 3:
            try:
                with st.spinner("🎭 Critic is reviewing the game… this may take a moment"):
                    cc_picks = _get_critics_choice(
                        active_lb,
                        special_instructions,
                        podium_names=[p["name"] for p in podium_players],
                        girl_names=girl_names,
                    )
                cc_quips = random.sample(_ALL_QUIPS, min(3, len(_ALL_QUIPS)))
                new_state = copy.deepcopy(_get())
                new_state["critics_choice"] = {"picks": cc_picks, "quips": cc_quips}
                with st.spinner("Saving…"):
                    _put(new_state)
                st.rerun()
            except Exception as _cc_err:
                st.error(f"Critic's analysis failed: {_cc_err}")
        else:
            st.warning("Need at least 3 players with games played.")

    cc_shared = state.get("critics_choice")
    if cc_shared and cc_shared.get("picks"):
        cc_picks  = cc_shared["picks"]
        cc_quips  = cc_shared.get("quips") or random.sample(_ALL_QUIPS, 3)
        _cc_bg     = ["#0D1526", "#130D26", "#1A0D1A"]
        _cc_border = ["#29B6F6", "#CE93D8", "#F48FB1"]
        _cc_medal_svgs = [
            _medal_svg("#FFD700", "#B8860B"),
            _medal_svg("#D8D8D8", "#888888"),
            _medal_svg("#CD7F32", "#8B4513"),
        ]
        cc_cols = st.columns(3)
        for col, medal_svg, pick, bg, border, quip_tmpl in zip(
            cc_cols, _cc_medal_svgs, cc_picks, _cc_bg, _cc_border, cc_quips
        ):
            quip = quip_tmpl.format(name=pick["name"])
            with col:
                st.markdown(
                    f'<div style="background:{bg};border-top:4px solid {border};'
                    f'border-radius:12px;padding:0.9rem 0.6rem;text-align:center;">'
                    f'<div style="line-height:1;margin-bottom:0.1rem">{medal_svg}</div>'
                    f'<div style="font-weight:700;font-size:0.95rem;margin-top:0.4rem;color:#E8EAF0;">'
                    f'{pick["name"]}</div>'
                    f'<div style="font-size:0.72rem;color:#aaa;margin-top:0.45rem;'
                    f'font-style:italic;line-height:1.4">{quip}</div>'
                    f'<div style="border-top:1px solid #2a2a3a;margin-top:0.55rem;'
                    f'padding-top:0.45rem;font-size:0.62rem;color:#90A4AE;'
                    f'line-height:1.5;text-align:left;">'
                    f'🎭 {pick["reason"]}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

    st.divider()

    # ── Table ────────────────────────────────────────────────────────────────
    rows = [
        {
            "#":      p["rank"],
            "Player": p["name"],
            "Pts":    p["games_won"] * 2,
            "W":      p["games_won"],
            "L":      p["games_lost"],
            "For":    p["points_gained"],
            "Agst":   p["points_conceded"],
            "Net":    p["net_points"],
        }
        for p in lb
    ]
    st.dataframe(
        pd.DataFrame(rows),
        use_container_width=True,
        hide_index=True,
        column_config={"Net": st.column_config.NumberColumn("Net", format="%+d")},
    )

    # ── Publish / Reset ───────────────────────────────────────────────────────
    st.divider()
    pub_col, rst_all_col, rst1_col = st.columns(3)
    with pub_col:
        if st.button("📤 Publish", type="primary", use_container_width=True):
            st.session_state.show_pub_pw   = True
            st.session_state.show_rst_pw   = False
            st.session_state.show_rst1_pw  = False
            st.session_state.show_rst1_select = False
            st.rerun()
    with rst_all_col:
        if st.button("🗑 Reset All", use_container_width=True):
            st.session_state.show_rst_pw   = True
            st.session_state.show_pub_pw   = False
            st.session_state.show_rst1_pw  = False
            st.session_state.show_rst1_select = False
            st.rerun()
    with rst1_col:
        if st.button("✂️ Reset 1 Game", use_container_width=True):
            st.session_state.show_rst1_pw  = True
            st.session_state.show_pub_pw   = False
            st.session_state.show_rst_pw   = False
            st.session_state.show_rst1_select = False
            st.rerun()

    # ── Publish password prompt ───────────────────────────────────────────────
    if st.session_state.show_pub_pw:
        st.caption("Enter password to publish leaderboard")
        pp_col, pg_col = st.columns([5, 2])
        with pp_col:
            pub_pw = st.text_input(
                "pub_pw", type="password", placeholder="Password…",
                label_visibility="collapsed", key="pub_pw_field",
            )
        with pg_col:
            if st.button("Confirm", key="btn_pub_confirm",
                         type="primary", use_container_width=True):
                if pub_pw == _admin_password():
                    st.session_state.show_pub_pw = False
                    _publish_leaderboard(lb)
                    st.success("✅ Leaderboard published!")
                else:
                    st.error("Incorrect password.")
                st.rerun()

    # ── Reset All — password then type-to-confirm ─────────────────────────────
    if st.session_state.show_rst_pw and not st.session_state.show_rst_confirm:
        st.caption("Enter password to reset all previous games")
        rp_col, rg_col = st.columns([5, 2])
        with rp_col:
            rst_pw = st.text_input(
                "rst_pw", type="password", placeholder="Password…",
                label_visibility="collapsed", key="rst_pw_field",
            )
        with rg_col:
            if st.button("Confirm", key="btn_rst_confirm",
                         type="primary", use_container_width=True):
                if rst_pw == _admin_password():
                    st.session_state.show_rst_pw      = False
                    st.session_state.show_rst_confirm = True
                else:
                    st.error("Incorrect password.")
                st.rerun()

    if st.session_state.show_rst_confirm:
        st.warning("This will delete **all** published game history. This cannot be undone.")
        st.caption('Type **reset** below to confirm')
        rc_col, rcc_col, rcancel_col = st.columns([4, 2, 2])
        with rc_col:
            confirm_text = st.text_input(
                "confirm_reset", placeholder='type "reset"',
                label_visibility="collapsed", key="rst_confirm_field",
            )
        with rcc_col:
            if st.button("Delete All", key="btn_rst_delete",
                         type="primary", use_container_width=True):
                if confirm_text.strip() == "reset":
                    st.session_state.show_rst_confirm = False
                    pass  # placeholder — functionality to be added later
                    st.info("Reset All functionality coming soon.")
                else:
                    st.error('Please type "reset" exactly.')
                st.rerun()
        with rcancel_col:
            if st.button("Cancel", key="btn_rst_cancel", use_container_width=True):
                st.session_state.show_rst_confirm = False
                st.session_state.show_rst_pw      = False
                st.rerun()

    # ── Reset 1 Game — password then selection ────────────────────────────────
    if st.session_state.show_rst1_pw and not st.session_state.show_rst1_select:
        st.caption("Enter password to delete a published entry")
        r1p_col, r1g_col = st.columns([5, 2])
        with r1p_col:
            rst1_pw = st.text_input(
                "rst1_pw", type="password", placeholder="Password…",
                label_visibility="collapsed", key="rst1_pw_field",
            )
        with r1g_col:
            if st.button("Confirm", key="btn_rst1_confirm",
                         type="primary", use_container_width=True):
                if rst1_pw == _admin_password():
                    st.session_state.show_rst1_pw     = False
                    st.session_state.show_rst1_select = True
                else:
                    st.error("Incorrect password.")
                st.rerun()

    if st.session_state.show_rst1_select:
        _pub2 = _github_load_file("data/published_results.json", default={"results": []})
        _all_results = (_pub2.get("results", []) if isinstance(_pub2, dict) else [])
        _recent = sorted(_all_results, key=lambda x: x.get("timestamp", ""), reverse=True)[:5]

        if not _recent:
            st.info("No published entries found.")
            st.session_state.show_rst1_select = False
        else:
            st.caption("Select entries to delete:")
            selected_idxs = []
            for i, entry in enumerate(_recent):
                ts = entry.get("timestamp", f"Entry {i + 1}")
                if st.checkbox(ts, key=f"rst1_chk_{i}"):
                    selected_idxs.append(i)

            del_col, cancel_col = st.columns(2)
            with del_col:
                if st.button("🗑 Delete Selected", type="primary", use_container_width=True):
                    if selected_idxs:
                        st.session_state.rst1_pending_delete = selected_idxs
                        st.session_state.show_rst1_confirm   = True
                    else:
                        st.warning("No entries selected.")
                    st.rerun()
            with cancel_col:
                if st.button("Cancel", use_container_width=True):
                    st.session_state.show_rst1_select  = False
                    st.session_state.show_rst1_confirm = False
                    st.rerun()

        if st.session_state.show_rst1_confirm:
            st.warning(f"About to delete **{len(st.session_state.rst1_pending_delete)}** entry/entries. This cannot be undone.")
            st.caption('Type **reset** below to confirm')
            r1c_col, r1cc_col = st.columns([5, 2])
            with r1c_col:
                rst1_confirm_text = st.text_input(
                    "rst1_confirm", placeholder='type "reset"',
                    label_visibility="collapsed", key="rst1_confirm_field",
                )
            with r1cc_col:
                if st.button("Delete", key="btn_rst1_delete",
                             type="primary", use_container_width=True):
                    if rst1_confirm_text.strip() == "reset":
                        timestamps_to_delete = {
                            _recent[i]["timestamp"]
                            for i in st.session_state.rst1_pending_delete
                        }
                        _all_results[:] = [
                            r for r in _all_results
                            if r.get("timestamp") not in timestamps_to_delete
                        ]
                        _github_save_file(
                            "data/published_results.json",
                            {"results": _all_results},
                            "leaderboard: delete published entries",
                        )
                        for i in st.session_state.rst1_pending_delete:
                            st.session_state.pop(f"rst1_chk_{i}", None)
                        st.session_state.show_rst1_select  = False
                        st.session_state.show_rst1_confirm = False
                        st.session_state.rst1_pending_delete = []
                        st.success(f"✅ Deleted {len(timestamps_to_delete)} entry/entries.")
                    else:
                        st.error('Please type "reset" exactly.')
                    st.rerun()



# ===========================================================================
# Page: Welcome / Phone Gate
# ===========================================================================

def show_welcome() -> None:
    st.markdown(
        '<div class="page-header" style="text-align:center;">'
        '<h1>🎾 Sports Leaderboard</h1>'
        '<p>AI Powered Tournament Scheduling : Doubles</p>'
        '</div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        '<p style="text-align:center;color:#aaa;margin:1.5rem 0 0.5rem;">'
        'Enter your mobile number to continue</p>',
        unsafe_allow_html=True,
    )

    # ── Phone input ───────────────────────────────────────────────────────────
    phone = st.text_input(
        "Mobile number", max_chars=10,
        placeholder="10-digit mobile number",
        label_visibility="collapsed",
        key="welcome_phone",
    )

    if st.button("Continue →", type="primary", use_container_width=True):
        p = (phone or "").strip()
        if len(p) != 10 or not p.isdigit():
            st.error("Please enter a valid 10-digit number.")
        else:
            users = _get_users()
            if p in users.get("allowed_phones", []):
                st.session_state.phone_verified = True
                st.session_state.verified_phone = p
                st.query_params["lb_auth"] = p
                st.session_state._ls_write_phone = p
                st.rerun()
            else:
                st.error("This number is not registered for this tournament.")

    # ── Admin section ─────────────────────────────────────────────────────────
    st.markdown('<div style="height:3rem"></div>', unsafe_allow_html=True)

    st.markdown(
        '<div style="text-align:center;margin-top:2rem;padding:0.75rem 1rem;'
        'border-radius:8px;background:#1A1F2E;color:#666;font-size:0.72rem;line-height:1.6">'
        'All tournament data — player names, schedules, and scores — is stored only for '
        'the duration of this tournament and will be cleared once the session is reset.'
        '</div>',
        unsafe_allow_html=True,
    )

    _, admin_col = st.columns([9, 1])
    with admin_col:
        st.markdown(
            """
            <style>
            .admin-dots-anchor + div,
            .admin-dots-anchor + div > div,
            .admin-dots-anchor + div > div > div { background:transparent!important;
                border:none!important; box-shadow:none!important; }
            .admin-dots-anchor + div button,
            .admin-dots-anchor + div button:hover,
            .admin-dots-anchor + div button:focus,
            .admin-dots-anchor + div button:active {
                min-height:22px!important; height:22px!important;
                width:auto!important; padding:0 0.4rem!important;
                font-size:0.6rem!important; background:transparent!important;
                border:none!important; outline:none!important;
                box-shadow:none!important; color:#2e2e3e!important;
                letter-spacing:0.2em; }
            .admin-dots-anchor + div button:hover { color:#555!important; }
            </style>
            <div class="admin-dots-anchor"></div>
            """,
            unsafe_allow_html=True,
        )
        if st.button("· · ·", key="btn_admin_trigger", use_container_width=False):
            st.session_state.show_admin_pw = not st.session_state.show_admin_pw
            st.session_state.show_admin_panel = False
            st.rerun()

    # ── Admin password prompt ─────────────────────────────────────────────────
    if st.session_state.show_admin_pw and not st.session_state.show_admin_panel:
        ap_col, ago_col = st.columns([5, 2])
        with ap_col:
            admin_pw = st.text_input(
                "apw", type="password", placeholder="Admin password…",
                label_visibility="collapsed", key="admin_pw_field",
            )
        with ago_col:
            if st.button("Unlock", key="btn_admin_unlock",
                         type="primary", use_container_width=True):
                if admin_pw == _admin_password():
                    st.session_state.show_admin_panel = True
                    st.session_state.show_admin_pw    = False
                    if "admin_pw_field" in st.session_state:
                        del st.session_state["admin_pw_field"]
                else:
                    st.error("Incorrect password.")
                st.rerun()

    # ── Admin panel ───────────────────────────────────────────────────────────
    if st.session_state.show_admin_panel:
        st.divider()
        st.markdown(
            '<div class="section-label">Manage registered numbers</div>',
            unsafe_allow_html=True,
        )
        users = _get_users()
        phones: List[str] = users.get("allowed_phones", [])

        # Current list
        if phones:
            for ph in phones:
                c_ph, c_del = st.columns([5, 1])
                c_ph.markdown(
                    f'<span style="font-size:0.95rem;letter-spacing:0.05em">'
                    f'{ph[:5]}·····</span>',
                    unsafe_allow_html=True,
                )
                if c_del.button("✕", key=f"del_{ph}", use_container_width=True):
                    phones.remove(ph)
                    with st.spinner("Saving…"):
                        _put_users({"allowed_phones": phones})
                    st.rerun()
        else:
            st.caption("No numbers registered yet.")

        # Add new number
        st.markdown('<div style="height:0.4rem"></div>', unsafe_allow_html=True)
        n_col, a_col = st.columns([5, 2])
        with n_col:
            new_ph = st.text_input(
                "new_ph", max_chars=10, placeholder="Add 10-digit number",
                label_visibility="collapsed",
                key=f"new_phone_field_{st.session_state.phone_add_counter}",
            )
        with a_col:
            if st.button("Add", key="btn_add_phone",
                         type="primary", use_container_width=True):
                p = (new_ph or "").strip()
                if len(p) == 10 and p.isdigit():
                    if p not in phones:
                        phones.append(p)
                        with st.spinner("Saving…"):
                            _put_users({"allowed_phones": phones})
                        st.session_state.phone_add_counter += 1
                        st.success(f"Added {p[:5]}·····")
                    else:
                        st.warning("Already registered.")
                else:
                    st.error("Enter a valid 10-digit number.")
                st.rerun()

        if st.button("Close admin", key="btn_admin_close", use_container_width=True):
            st.session_state.show_admin_panel = False
            st.rerun()


# ===========================================================================
# Router
# ===========================================================================

if not st.session_state.phone_verified:
    show_welcome()
else:
    # Write phone to localStorage once right after login
    if st.session_state.get("_ls_write_phone"):
        _lsp = st.session_state.pop("_ls_write_phone")
        _lst = int(_time.time())
        st.components.v1.html(
            f'<script>try{{window.localStorage.setItem("lb_auth",'
            f'JSON.stringify({{p:"{_lsp}",t:{_lst}}}))}}catch(e){{}}</script>',
            height=0,
        )

    _page = st.session_state.page
    _nav(_page)
    st.markdown("---")
    if _page == "setup":
        show_setup()
    elif _page == "leaderboard":
        show_leaderboard()
    elif _page == "alltime":
        show_all_time_leaderboard()
    elif _page.startswith("court"):
        try:
            show_court(int(_page[5:]))
        except ValueError:
            show_setup()
    else:
        show_setup()
